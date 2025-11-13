# WordLibrary.py
# Reusable Robot Framework library for generating Word (.docx) reports
# Features: Create/Open, Headings, Text (bold/italic/align), Images with Captions, Save

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Correct import
from robot.libraries.BuiltIn import BuiltIn


class WordLibrary:
    """
    Standalone Word document generator for Robot Framework.
    - One document per test suite (ROBOT_LIBRARY_SCOPE = 'TEST SUITE')
    - Supports: add headings, add text, images with captions.
    - Reusable in any .robot file
    """
    ROBOT_LIBRARY_SCOPE = 'TEST SUITE'

    def __init__(self):
        self.doc = None
        self.file_path = None
        self._builtin = BuiltIn()  # For logging to Robot log.html

    # ==============================================================
    # Document creation, opening and saving
    # ==============================================================
    # Keyword: Create Word Document. 1 Argument- file_path
    def create_word_document(self, file_path):
        self.doc = Document()
        self.file_path = file_path
        self._builtin.log(f"Created new Word document: {file_path}")

    def open_word_document(self, file_path):
        self.doc = Document(file_path)
        self.file_path = file_path
        self._builtin.log(f"Opened existing Word document: {file_path}")

    # Keyword: Save Document . No need of arguments as it will call for the arg in create mtd.
    def save_document(self):
        self.doc.save(self.file_path)
        self._builtin.log(f"Document saved: {self.file_path}")

    # ==============================================================
    # text & Paragraphs
    # ==============================================================

    def add_heading(self, text, level=1):
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text="", style=None):
        self.doc.add_paragraph(text, style=style)

    def add_text(self, text, bold=False, italic=False, size=None, align=None):
        """
        Add formatted text in a new paragraph.Below are the options.
        - bold: True/False
        - italic: True/False
        - size: font size in Pt (e.g. 12)
        - align: 'left', 'center', 'right', 'justify'
        """
        p = self.doc.add_paragraph()
        run = p.add_run(text)

        run.bold = bold
        run.italic = italic

        if size is not None:
            run.font.size = Pt(size)

        if align:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            p.alignment = align_map.get(align.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    # ==============================================================
    # IMAGES WITH CAPTIONS
    # ==============================================================
    # Keyword: Add Image With Caption. 2 arguments -image path and caption
    def add_image_with_caption(self, image_path, caption, width_inch=5):
        """
        Add an image with an italic, centered caption below it.
        - width_inch: image width in inches (default: 5)
        """
        if not os.path.exists(image_path):
            self._builtin.log(f"Image not found: {image_path}", level='WARN')
            return

        # Caption (italic, centered)
        cap_para = self.doc.add_paragraph()
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Image (centered)
        img_para = self.doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.add_run().add_picture(image_path, width=Inches(width_inch))

        # Add space
        self.doc.add_paragraph()

    # ==============================================================
    # UTILITIES
    # ==============================================================

    def add_page_break(self):
        """Insert a page break."""
        self.doc.add_page_break()